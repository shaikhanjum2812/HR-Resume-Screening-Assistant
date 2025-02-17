from docx import Document
import io

class DOCXProcessor:
    def extract_text(self, uploaded_file):
        try:
            # Create a file-like object from the uploaded file
            docx_file = io.BytesIO(uploaded_file.read())
            
            # Load the document
            doc = Document(docx_file)
            
            # Extract text from all paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to process DOCX: {e}")
